"""
Section Renderer
==================
Chia markdown thành sections nhỏ, render từng section riêng,
sau đó merge lại thành file DOCX hoàn chỉnh.

Đảm bảo xử lý được file luận văn/báo cáo lớn với full scan content.
"""

import os
import re
from pathlib import Path
from typing import List, Dict, Tuple, Optional
from dataclasses import dataclass
from enum import Enum


class SectionType(Enum):
    """Loại section"""
    FRONTMATTER = "frontmatter"  # Tiêu đề, mục lục...
    CHAPTER = "chapter"          # Chương (H1)
    SECTION = "section"          # Mục (H2)
    SUBSECTION = "subsection"    # Mục con (H3+)
    SEPARATOR = "separator"      # ---


@dataclass
class Section:
    """Đại diện cho một section"""
    title: str
    content: str
    section_type: SectionType
    level: int  # 1 = H1, 2 = H2, etc.
    index: int  # Thứ tự trong document
    
    def __repr__(self):
        preview = self.content[:50].replace('\n', ' ')
        return f"Section({self.index}: {self.section_type.value} L{self.level} '{self.title}' - {preview}...)"


class SectionSplitter:
    """
    Chia markdown thành các sections dựa trên:
    - Horizontal rule: ---
    - H1 heading: #
    - H2 heading: ##
    - H3+ heading: ###, ####, etc.
    """
    
    def __init__(self, split_on_hr: bool = True, 
                 split_on_h1: bool = True,
                 split_on_h2: bool = True,
                 min_section_lines: int = 3):
        """
        Args:
            split_on_hr: Chia theo horizontal rule ---
            split_on_h1: Chia theo H1 #
            split_on_h2: Chia theo H2 ##
            min_section_lines: Số dòng tối thiểu mỗi section
        """
        self.split_on_hr = split_on_hr
        self.split_on_h1 = split_on_h1
        self.split_on_h2 = split_on_h2
        self.min_section_lines = min_section_lines
    
    def split(self, markdown_content: str) -> List[Section]:
        """
        Chia markdown thành danh sách sections
        
        Args:
            markdown_content: Nội dung markdown đầy đủ
        
        Returns:
            List các Section objects
        """
        lines = markdown_content.split('\n')
        sections = []
        current_section_lines = []
        current_title = "Introduction"
        current_type = SectionType.FRONTMATTER
        current_level = 0
        section_index = 0
        
        for i, line in enumerate(lines):
            is_split_point = False
            new_title = None
            new_type = None
            new_level = 0
            
            # Check horizontal rule
            if self.split_on_hr and re.match(r'^-{3,}\s*$', line.strip()):
                is_split_point = True
                new_title = f"Section {section_index + 1}"
                new_type = SectionType.SEPARATOR
                new_level = 0
            
            # Check H1
            elif self.split_on_h1 and re.match(r'^#\s+(.+)$', line):
                match = re.match(r'^#\s+(.+)$', line)
                is_split_point = True
                new_title = match.group(1).strip()
                new_type = SectionType.CHAPTER
                new_level = 1
            
            # Check H2
            elif self.split_on_h2 and re.match(r'^##\s+(.+)$', line):
                match = re.match(r'^##\s+(.+)$', line)
                is_split_point = True
                new_title = match.group(1).strip()
                new_type = SectionType.SECTION
                new_level = 2
            
            if is_split_point:
                # Save previous section nếu có content
                if current_section_lines:
                    content = '\n'.join(current_section_lines)
                    if content.strip():
                        sections.append(Section(
                            title=current_title,
                            content=content,
                            section_type=current_type,
                            level=current_level,
                            index=section_index
                        ))
                        section_index += 1
                
                # Start new section
                current_section_lines = [line]
                current_title = new_title
                current_type = new_type
                current_level = new_level
            else:
                current_section_lines.append(line)
        
        # Save last section
        if current_section_lines:
            content = '\n'.join(current_section_lines)
            if content.strip():
                sections.append(Section(
                    title=current_title,
                    content=content,
                    section_type=current_type,
                    level=current_level,
                    index=section_index
                ))
        
        return sections
    
    def split_file(self, file_path: str) -> List[Section]:
        """Chia file markdown"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        return self.split(content)


class SectionRenderer:
    """
    Render từng section thành DOCX riêng, sau đó merge lại
    """
    
    def __init__(self, output_dir: str = "temp_sections"):
        self.output_dir = output_dir
        self.splitter = SectionSplitter()
    
    def render_sections(self, markdown_path: str, 
                        final_output: str = None) -> str:
        """
        Render markdown thành DOCX qua từng section
        
        Args:
            markdown_path: Path đến file markdown
            final_output: Path output cuối cùng
        
        Returns:
            Path đến file DOCX hoàn chỉnh
        """
        # 1. Split markdown
        sections = self.splitter.split_file(markdown_path)
        print(f"📄 Split into {len(sections)} sections")
        
        for sec in sections:
            print(f"  {sec.index}: [{sec.section_type.value}] {sec.title}")
        
        # 2. Create temp directory
        temp_dir = Path(self.output_dir)
        temp_dir.mkdir(parents=True, exist_ok=True)
        
        # 3. Save and render each section
        section_files = []
        
        for section in sections:
            # Save section markdown
            section_md_path = temp_dir / f"section_{section.index:03d}.md"
            with open(section_md_path, 'w', encoding='utf-8') as f:
                f.write(section.content)
            
            # Render to DOCX
            section_docx_path = temp_dir / f"section_{section.index:03d}.docx"
            
            try:
                self._render_single_section(
                    str(section_md_path), 
                    str(section_docx_path),
                    section
                )
                section_files.append(str(section_docx_path))
                print(f"  ✓ Rendered: section_{section.index:03d}.docx")
            except Exception as e:
                print(f"  ⚠ Error rendering section {section.index}: {e}")
        
        # 4. Merge all sections
        if not final_output:
            base_name = Path(markdown_path).stem
            final_output = temp_dir.parent / f"{base_name}_rendered.docx"
        
        merged_path = self._merge_sections(section_files, str(final_output))
        print(f"\n✅ Final document: {merged_path}")
        
        return merged_path
    
    def _render_single_section(self, md_path: str, docx_path: str, 
                                section: Section):
        """Render một section với markdown cleaning"""
        try:
            from function2.templates.converters.md_to_docx import MarkdownToDocx
            
            converter = MarkdownToDocx()
            converter.convert(md_path, docx_path)
            
            # Verify file was created
            if not os.path.exists(docx_path):
                raise Exception(f"File not created: {docx_path}")
        except Exception as e:
            # Fallback: create simple docx
            print(f"    Using fallback for section {section.index}")
            self._create_simple_docx(md_path, docx_path)
    
    def _create_simple_docx(self, md_path: str, docx_path: str):
        """Create simple docx without complex conversion"""
        from docx import Document
        from docx.shared import Pt, Cm
        from src.templates.markdown_cleaner import MarkdownCleaner
        
        with open(md_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        doc = Document()
        cleaner = MarkdownCleaner()
        
        for line in content.split('\n'):
            stripped = line.strip()
            if not stripped:
                continue
            
            # Handle headings
            if stripped.startswith('#'):
                level = len(stripped.split()[0])
                text = stripped.lstrip('#').strip()
                para = doc.add_paragraph()
                run = para.add_run(text)
                run.font.bold = True
                run.font.size = Pt(16 - level)
            else:
                # Regular paragraph with inline formatting
                para = doc.add_paragraph()
                cleaner.apply_to_paragraph(para, stripped)
        
        doc.save(docx_path)
    
    def _merge_sections(self, section_files: List[str], output_path: str) -> str:
        """Merge tất cả section DOCX thành một file"""
        # Filter only existing files
        existing_files = [f for f in section_files if os.path.exists(f)]
        
        print(f"  📁 Merging {len(existing_files)} files...")
        
        if not existing_files:
            raise ValueError("No section files to merge")
        
        # If only one file, just copy it
        if len(existing_files) == 1:
            import shutil
            shutil.copy(existing_files[0], output_path)
            return output_path
        
        # Merge multiple files - use seamless (no page breaks)
        try:
            from function2.templates.converters.docx_merger import merge_docx_seamless
            return merge_docx_seamless(existing_files, output_path)
        except ImportError:
            # Fallback merge
            return self._simple_merge(existing_files, output_path)


class MarkdownScanner:
    """
    Scan toàn bộ nội dung markdown để đảm bảo không bị mất data
    """
    
    def __init__(self):
        pass
    
    def scan(self, markdown_content: str) -> Dict:
        """
        Scan và thống kê nội dung markdown
        
        Returns:
            Dict với các thống kê
        """
        lines = markdown_content.split('\n')
        
        stats = {
            "total_lines": len(lines),
            "total_chars": len(markdown_content),
            "headings": [],
            "paragraphs": 0,
            "lists": 0,
            "code_blocks": 0,
            "images": 0,
            "links": 0,
            "tables": 0,
            "horizontal_rules": 0,
            "bold_count": 0,
            "italic_count": 0,
        }
        
        in_code_block = False
        
        for line in lines:
            stripped = line.strip()
            
            # Code blocks
            if stripped.startswith('```'):
                if in_code_block:
                    in_code_block = False
                else:
                    in_code_block = True
                    stats["code_blocks"] += 1
                continue
            
            if in_code_block:
                continue
            
            # Headings
            heading_match = re.match(r'^(#{1,6})\s+(.+)$', stripped)
            if heading_match:
                level = len(heading_match.group(1))
                title = heading_match.group(2)
                stats["headings"].append({
                    "level": level,
                    "title": title,
                    "line": lines.index(line) + 1
                })
                continue
            
            # Horizontal rules
            if re.match(r'^-{3,}\s*$', stripped) or re.match(r'^\*{3,}\s*$', stripped):
                stats["horizontal_rules"] += 1
                continue
            
            # Lists
            if re.match(r'^[-*+]\s+', stripped) or re.match(r'^\d+\.\s+', stripped):
                stats["lists"] += 1
                continue
            
            # Tables
            if '|' in stripped and '-' in stripped:
                stats["tables"] += 1
                continue
            
            # Regular paragraphs
            if stripped:
                stats["paragraphs"] += 1
        
        # Count inline elements
        stats["bold_count"] = len(re.findall(r'\*\*[^*]+\*\*', markdown_content))
        stats["italic_count"] = len(re.findall(r'\*[^*]+\*', markdown_content))
        stats["images"] = len(re.findall(r'!\[.*?\]\(.*?\)', markdown_content))
        stats["links"] = len(re.findall(r'\[.*?\]\(.*?\)', markdown_content))
        
        return stats
    
    def scan_file(self, file_path: str) -> Dict:
        """Scan file markdown"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        return self.scan(content)
    
    def verify_completeness(self, original_content: str, 
                            sections: List[Section]) -> Tuple[bool, List[str]]:
        """
        Verify rằng tất cả nội dung đã được bao gồm trong sections
        
        Returns:
            (is_complete, list of missing parts)
        """
        original_lines = set(line.strip() for line in original_content.split('\n') if line.strip())
        
        section_lines = set()
        for section in sections:
            for line in section.content.split('\n'):
                if line.strip():
                    section_lines.add(line.strip())
        
        missing = original_lines - section_lines
        
        return len(missing) == 0, list(missing)[:10]  # Return first 10 missing
    
    def print_report(self, stats: Dict):
        """In báo cáo scan"""
        print("\n📊 Markdown Scan Report")
        print("=" * 40)
        print(f"📄 Total lines: {stats['total_lines']}")
        print(f"📝 Total chars: {stats['total_chars']}")
        print(f"📑 Headings: {len(stats['headings'])}")
        print(f"📃 Paragraphs: {stats['paragraphs']}")
        print(f"📋 Lists: {stats['lists']}")
        print(f"💻 Code blocks: {stats['code_blocks']}")
        print(f"🖼️ Images: {stats['images']}")
        print(f"🔗 Links: {stats['links']}")
        print(f"📊 Tables: {stats['tables']}")
        print(f"➖ Horizontal rules: {stats['horizontal_rules']}")
        print(f"**Bold**: {stats['bold_count']}")
        print(f"*Italic*: {stats['italic_count']}")
        
        if stats['headings']:
            print("\n📑 Heading Structure:")
            for h in stats['headings']:
                indent = "  " * (h['level'] - 1)
                print(f"  {indent}{'#' * h['level']} {h['title']}")


def render_with_sections(markdown_path: str, output_path: str = None) -> str:
    """
    Hàm tiện ích để render markdown qua section splitting
    
    Args:
        markdown_path: Path đến file markdown
        output_path: Path output (optional)
    
    Returns:
        Path đến file DOCX
    """
    # 1. Scan first
    scanner = MarkdownScanner()
    stats = scanner.scan_file(markdown_path)
    scanner.print_report(stats)
    
    # 2. Split và verify
    splitter = SectionSplitter()
    sections = splitter.split_file(markdown_path)
    
    with open(markdown_path, 'r', encoding='utf-8') as f:
        original = f.read()
    
    is_complete, missing = scanner.verify_completeness(original, sections)
    
    if not is_complete:
        print(f"\n⚠ Warning: {len(missing)} lines may be missing")
        for line in missing[:5]:
            print(f"  - {line[:50]}...")
    else:
        print("\n✅ All content accounted for")
    
    # 3. Render
    renderer = SectionRenderer(
        output_dir=str(Path(markdown_path).parent / "temp_sections")
    )
    
    return renderer.render_sections(markdown_path, output_path)


if __name__ == "__main__":
    import sys
    
    if len(sys.argv) > 1:
        md_path = sys.argv[1]
        output = sys.argv[2] if len(sys.argv) > 2 else None
        
        result = render_with_sections(md_path, output)
        print(f"\n✅ Output: {result}")
    else:
        # Demo
        sample = """# Chương 1: Giới thiệu

Đây là phần giới thiệu với **bold text**.

---

## 1.1 Đặt vấn đề

Nội dung đặt vấn đề với *italic*.

## 1.2 Mục tiêu

- Mục tiêu 1
- Mục tiêu 2

---

# Chương 2: Cơ sở lý thuyết

**Kết quả mong đợi:** Báo cáo hoàn chỉnh.

## 2.1 Lý thuyết A

Nội dung lý thuyết...
"""
        
        splitter = SectionSplitter()
        sections = splitter.split(sample)
        
        print(f"Split into {len(sections)} sections:\n")
        for sec in sections:
            print(f"[{sec.index}] {sec.section_type.value} - {sec.title}")
            print(f"    Content: {sec.content[:60]}...")
            print()
        
        scanner = MarkdownScanner()
        stats = scanner.scan(sample)
        scanner.print_report(stats)
