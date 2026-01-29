"""
DOCX Merger - Ghép nhiều file Word thành 1 file duy nhất
=========================================================
Sử dụng python-docx để merge các sections thành 1 document hoàn chỉnh

Features:
- Giữ nguyên styles, fonts, formatting
- Đánh số trang liên tục
- Tạo Table of Contents tự động
- Thêm page breaks giữa các sections
"""

import os
from pathlib import Path
from typing import List, Optional, Dict
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
import copy


class DocxMerger:
    """Merge nhiều file DOCX thành 1 file duy nhất"""
    
    def __init__(self, output_path: str = None):
        self.output_path = output_path
        self.merged_doc = Document()
        self._setup_document()
    
    def _setup_document(self):
        """Thiết lập document mới với styles chuẩn"""
        # Set default font
        style = self.merged_doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)
        
        # Set margins theo NĐ30/2020
        for section in self.merged_doc.sections:
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(3.0)
            section.right_margin = Cm(1.5)
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)
    
    def _copy_element(self, element):
        """Deep copy một element XML"""
        return copy.deepcopy(element)
    
    def _add_page_break(self):
        """Thêm page break giữa các sections"""
        self.merged_doc.add_page_break()
    
    def add_document(self, doc_path: str, add_page_break: bool = True):
        """
        Thêm nội dung từ 1 file DOCX vào document chính
        
        Args:
            doc_path: Đường dẫn đến file DOCX
            add_page_break: Có thêm page break trước không
        """
        if not os.path.exists(doc_path):
            raise FileNotFoundError(f"File not found: {doc_path}")
        
        # Load source document
        source_doc = Document(doc_path)
        
        # Add page break before content (except first document)
        if add_page_break and len(self.merged_doc.paragraphs) > 0:
            self._add_page_break()
        
        # Copy all paragraphs
        for para in source_doc.paragraphs:
            # Create new paragraph in merged doc
            new_para = self.merged_doc.add_paragraph()
            new_para.style = para.style
            new_para.alignment = para.alignment
            
            # Copy paragraph format
            if para.paragraph_format.first_line_indent:
                new_para.paragraph_format.first_line_indent = para.paragraph_format.first_line_indent
            
            # Copy all runs (text with formatting)
            for run in para.runs:
                new_run = new_para.add_run(run.text)
                # Copy font formatting
                new_run.font.name = run.font.name or 'Times New Roman'
                new_run.font.size = run.font.size or Pt(14)
                new_run.font.bold = run.font.bold
                new_run.font.italic = run.font.italic
                new_run.font.underline = run.font.underline
        
        # Copy all tables
        for table in source_doc.tables:
            self._copy_table(table)
        
        return self
    
    def _copy_table(self, source_table):
        """Copy một table từ source document"""
        # Get table dimensions
        rows = len(source_table.rows)
        cols = len(source_table.columns)
        
        # Create new table
        new_table = self.merged_doc.add_table(rows=rows, cols=cols)
        new_table.style = source_table.style
        
        # Copy cell contents
        for i, row in enumerate(source_table.rows):
            for j, cell in enumerate(row.cells):
                new_cell = new_table.cell(i, j)
                # Copy text
                new_cell.text = cell.text
                # Copy paragraph formatting
                for para in new_cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        return new_table
    
    def merge_files(self, file_paths: List[str], add_page_breaks: bool = True) -> 'DocxMerger':
        """
        Merge nhiều files DOCX
        
        Args:
            file_paths: List đường dẫn đến các file DOCX
            add_page_breaks: Có thêm page break giữa các files không
        
        Returns:
            self để chain methods
        """
        for i, path in enumerate(file_paths):
            # First file doesn't need page break before
            add_break = add_page_breaks and i > 0
            self.add_document(path, add_page_break=add_break)
            print(f"✓ Merged: {os.path.basename(path)}")
        
        return self
    
    def merge_folder(self, folder_path: str, pattern: str = "*.docx", 
                     sort_by_name: bool = True) -> 'DocxMerger':
        """
        Merge tất cả file DOCX trong 1 folder
        
        Args:
            folder_path: Đường dẫn folder
            pattern: Pattern để filter files (default: *.docx)
            sort_by_name: Sắp xếp theo tên file
        
        Returns:
            self để chain methods
        """
        folder = Path(folder_path)
        if not folder.exists():
            raise FileNotFoundError(f"Folder not found: {folder_path}")
        
        # Get all matching files
        files = list(folder.glob(pattern))
        
        if sort_by_name:
            files = sorted(files, key=lambda x: x.name)
        
        if not files:
            print(f"⚠ No files matching '{pattern}' in {folder_path}")
            return self
        
        print(f"Found {len(files)} files to merge:")
        for f in files:
            print(f"  - {f.name}")
        
        # Merge all files
        self.merge_files([str(f) for f in files])
        
        return self
    
    def add_table_of_contents(self, title: str = "MỤC LỤC"):
        """
        Thêm Table of Contents placeholder
        (Cần update manually trong Word để hiển thị đầy đủ)
        """
        # Title
        toc_title = self.merged_doc.add_paragraph()
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = toc_title.add_run(title)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True
        
        # TOC field code
        paragraph = self.merged_doc.add_paragraph()
        run = paragraph.add_run()
        
        # Add TOC field
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)
        
        # Add instruction
        note = self.merged_doc.add_paragraph()
        note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = note.add_run("[Nhấn Ctrl+A → F9 để update mục lục]")
        run.font.size = Pt(11)
        run.font.italic = True
        
        self._add_page_break()
        
        return self
    
    def add_page_numbers(self):
        """Thêm đánh số trang (footer)"""
        for section in self.merged_doc.sections:
            footer = section.footer
            footer.is_linked_to_previous = False
            
            # Clear existing content
            footer.paragraphs[0].clear()
            
            # Add page number
            paragraph = footer.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run = paragraph.add_run()
            
            # Add PAGE field
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            
            instrText = OxmlElement('w:instrText')
            instrText.text = 'PAGE'
            
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            
            run._r.append(fldChar1)
            run._r.append(instrText)
            run._r.append(fldChar2)
        
        return self
    
    def save(self, output_path: str = None) -> str:
        """
        Lưu document đã merge
        
        Args:
            output_path: Đường dẫn output (optional, dùng self.output_path nếu không có)
        
        Returns:
            Đường dẫn file đã lưu
        """
        path = output_path or self.output_path
        if not path:
            raise ValueError("Output path is required")
        
        # Create directory if not exists
        os.makedirs(os.path.dirname(path), exist_ok=True)
        
        self.merged_doc.save(path)
        print(f"✅ Saved merged document: {path}")
        
        return path


def merge_docx_files(input_files: List[str], output_file: str, 
                     add_toc: bool = True, add_page_numbers: bool = True) -> str:
    """
    Hàm tiện ích để merge nhiều file DOCX
    
    Args:
        input_files: List các file DOCX cần merge
        output_file: Đường dẫn file output
        add_toc: Có thêm mục lục không
        add_page_numbers: Có đánh số trang không
    
    Returns:
        Đường dẫn file đã merge
    
    Example:
        >>> merge_docx_files(
        ...     ['section_001.docx', 'section_002.docx'],
        ...     'MERGED_document.docx'
        ... )
    """
    merger = DocxMerger(output_file)
    
    if add_toc:
        merger.add_table_of_contents()
    
    merger.merge_files(input_files)
    
    if add_page_numbers:
        merger.add_page_numbers()
    
    return merger.save()


def merge_docx_folder(folder_path: str, output_file: str,
                      pattern: str = "*.docx",
                      add_toc: bool = True,
                      add_page_numbers: bool = True) -> str:
    """
    Merge tất cả file DOCX trong folder
    
    Args:
        folder_path: Đường dẫn folder chứa files
        output_file: Đường dẫn file output
        pattern: Pattern filter (default: *.docx)
        add_toc: Có thêm mục lục không
        add_page_numbers: Có đánh số trang không
    
    Returns:
        Đường dẫn file đã merge
    
    Example:
        >>> merge_docx_folder(
        ...     'Segmentation/phase5_output/sections/',
        ...     'Segmentation/phase5_output/MERGED_document.docx'
        ... )
    """
    merger = DocxMerger(output_file)
    
    if add_toc:
        merger.add_table_of_contents()
    
    merger.merge_folder(folder_path, pattern)
    
    if add_page_numbers:
        merger.add_page_numbers()
    
    return merger.save()


if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description="Merge nhiều file DOCX thành 1")
    parser.add_argument('--folder', '-f', help='Folder chứa các file DOCX')
    parser.add_argument('--files', '-i', nargs='+', help='List các file DOCX')
    parser.add_argument('--output', '-o', required=True, help='File output')
    parser.add_argument('--no-toc', action='store_true', help='Không thêm mục lục')
    parser.add_argument('--no-page-numbers', action='store_true', help='Không đánh số trang')
    
    args = parser.parse_args()
    
    if args.folder:
        merge_docx_folder(
            args.folder, 
            args.output,
            add_toc=not args.no_toc,
            add_page_numbers=not args.no_page_numbers
        )
    elif args.files:
        merge_docx_files(
            args.files,
            args.output,
            add_toc=not args.no_toc,
            add_page_numbers=not args.no_page_numbers
        )
    else:
        print("Please specify --folder or --files")
        parser.print_help()
