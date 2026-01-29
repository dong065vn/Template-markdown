"""
Markdown to DOCX Converter
===========================
Chuyển đổi file Markdown thành DOCX với styles theo NĐ30/2020
Sử dụng Python templates để đảm bảo format 100% chuẩn
"""

import os
import re
from pathlib import Path
from typing import Dict, Optional, List
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


# Constants
FONT_NAME = "Times New Roman"
FONT_SIZE_BODY = 14
FONT_SIZE_HEADING1 = 14
FONT_SIZE_HEADING2 = 14
FONT_SIZE_HEADING3 = 14


class MarkdownToDocx:
    """Chuyển đổi Markdown sang DOCX với styles chuẩn NĐ30/2020"""
    
    def __init__(self, template_type: str = "thesis"):
        self.template_type = template_type
        self.document = Document()
        self._setup_document()
    
    def _setup_document(self):
        """Thiết lập document với styles chuẩn"""
        # Set margins
        for section in self.document.sections:
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(3.0)
            section.right_margin = Cm(1.5)
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)
        
        # Set default font
        style = self.document.styles['Normal']
        font = style.font
        font.name = FONT_NAME
        font.size = Pt(FONT_SIZE_BODY)
        
        # Set Vietnamese font
        rFonts = style.element.rPr.rFonts
        rFonts.set(qn('w:eastAsia'), FONT_NAME)
    
    def convert_file(self, md_path: str, output_path: str = None) -> str:
        """
        Chuyển đổi file Markdown sang DOCX
        
        Args:
            md_path: Đường dẫn file Markdown
            output_path: Đường dẫn file output (optional)
        
        Returns:
            Đường dẫn file DOCX đã tạo
        """
        if not os.path.exists(md_path):
            raise FileNotFoundError(f"File not found: {md_path}")
        
        with open(md_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        self.convert_content(md_content)
        
        if output_path is None:
            output_path = str(Path(md_path).with_suffix('.docx'))
        
        return self.save(output_path)
    
    def convert_content(self, md_content: str) -> 'MarkdownToDocx':
        """
        Chuyển đổi nội dung Markdown sang DOCX
        
        Args:
            md_content: Nội dung Markdown
        
        Returns:
            self để chain methods
        """
        lines = md_content.split('\n')
        current_list_type = None
        list_counter = 0
        
        for line in lines:
            line = line.rstrip()
            
            # Skip empty lines
            if not line.strip():
                continue
            
            # Heading 1: # Title
            if line.startswith('# '):
                self._add_heading1(line[2:].strip())
            
            # Heading 2: ## Title
            elif line.startswith('## '):
                self._add_heading2(line[3:].strip())
            
            # Heading 3: ### Title
            elif line.startswith('### '):
                self._add_heading3(line[4:].strip())
            
            # Heading 4: #### Title
            elif line.startswith('#### '):
                self._add_heading4(line[5:].strip())
            
            # Unordered list: - item or * item
            elif line.strip().startswith('- ') or line.strip().startswith('* '):
                item = line.strip()[2:].strip()
                self._add_list_item(item, ordered=False)
                current_list_type = 'unordered'
            
            # Ordered list: 1. item
            elif re.match(r'^\d+\.\s', line.strip()):
                item = re.sub(r'^\d+\.\s', '', line.strip())
                list_counter += 1
                self._add_list_item(item, ordered=True, number=list_counter)
                current_list_type = 'ordered'
            
            # Blockquote: > text
            elif line.strip().startswith('> '):
                quote = line.strip()[2:].strip()
                self._add_blockquote(quote)
            
            # Code block: ``` (skip for now)
            elif line.strip().startswith('```'):
                continue
            
            # Normal paragraph
            else:
                self._add_paragraph(line.strip())
                current_list_type = None
                list_counter = 0
        
        return self
    
    def _add_heading1(self, text: str):
        """Thêm Heading 1 (CHƯƠNG)"""
        para = self.document.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)
        
        run = para.add_run(text.upper())
        run.font.name = FONT_NAME
        run.font.size = Pt(FONT_SIZE_HEADING1)
        run.font.bold = True
    
    def _add_heading2(self, text: str):
        """Thêm Heading 2 (Mục lớn)"""
        para = self.document.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)
        
        run = para.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = Pt(FONT_SIZE_HEADING2)
        run.font.bold = True
    
    def _add_heading3(self, text: str):
        """Thêm Heading 3 (Mục nhỏ)"""
        para = self.document.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.first_line_indent = Cm(1.27)
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
        
        run = para.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = Pt(FONT_SIZE_HEADING3)
        run.font.bold = True
        run.font.italic = True
    
    def _add_heading4(self, text: str):
        """Thêm Heading 4"""
        para = self.document.add_paragraph()
        para.paragraph_format.first_line_indent = Cm(1.27)
        
        run = para.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = Pt(FONT_SIZE_BODY)
        run.font.bold = True
    
    def _add_paragraph(self, text: str):
        """Thêm đoạn văn thông thường"""
        # Parse inline formatting
        text = self._parse_inline_formatting(text)
        
        para = self.document.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.first_line_indent = Cm(1.27)
        para.paragraph_format.space_after = Pt(6)
        
        # Add text with inline formatting
        self._add_formatted_text(para, text)
    
    def _parse_inline_formatting(self, text: str) -> str:
        """Parse inline markdown formatting (để xử lý sau)"""
        return text
    
    def _add_formatted_text(self, para, text: str):
        """Thêm text với inline formatting (**bold**, *italic*, ***bold+italic***)"""
        # Pattern để tìm các phần **bold**, *italic*, ***bold+italic***
        # Thứ tự: bold+italic trước, rồi bold, rồi italic
        pattern = r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|([^*]+))'
        
        parts = re.findall(pattern, text)
        
        if not parts or text == ''.join(p[0] for p in parts if p[0]):
            # Fallback: không có formatting
            for match in re.finditer(pattern, text):
                full, bold_italic, bold, italic, plain = match.groups()
                
                if bold_italic:
                    run = para.add_run(bold_italic)
                    run.font.bold = True
                    run.font.italic = True
                elif bold:
                    run = para.add_run(bold)
                    run.font.bold = True
                elif italic:
                    run = para.add_run(italic)
                    run.font.italic = True
                elif plain:
                    run = para.add_run(plain)
                else:
                    continue
                
                run.font.name = FONT_NAME
                run.font.size = Pt(FONT_SIZE_BODY)
        else:
            # Simple fallback without formatting
            run = para.add_run(text)
            run.font.name = FONT_NAME
            run.font.size = Pt(FONT_SIZE_BODY)
    
    def _add_list_item(self, text: str, ordered: bool = False, number: int = 1):
        """Thêm list item"""
        para = self.document.add_paragraph()
        para.paragraph_format.left_indent = Cm(1.27)
        
        if ordered:
            prefix = f"{number}. "
        else:
            prefix = "• "
        
        run = para.add_run(prefix + text)
        run.font.name = FONT_NAME
        run.font.size = Pt(FONT_SIZE_BODY)
    
    def _add_blockquote(self, text: str):
        """Thêm blockquote"""
        para = self.document.add_paragraph()
        para.paragraph_format.left_indent = Cm(2.0)
        para.paragraph_format.right_indent = Cm(1.0)
        
        run = para.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = Pt(FONT_SIZE_BODY)
        run.font.italic = True
    
    def save(self, output_path: str) -> str:
        """Lưu document"""
        os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else '.', exist_ok=True)
        self.document.save(output_path)
        print(f"✅ Saved: {output_path}")
        return output_path


def convert_md_to_docx(md_path: str, output_path: str = None, 
                       template_type: str = "thesis") -> str:
    """
    Hàm tiện ích để chuyển đổi Markdown sang DOCX
    
    Args:
        md_path: Đường dẫn file Markdown
        output_path: Đường dẫn output (optional)
        template_type: Loại template (thesis, report, official)
    
    Returns:
        Đường dẫn file DOCX
    
    Example:
        >>> convert_md_to_docx('content_001.md', 'section_001.docx')
    """
    converter = MarkdownToDocx(template_type)
    return converter.convert_file(md_path, output_path)


def convert_folder(input_folder: str, output_folder: str,
                   pattern: str = "*.md") -> List[str]:
    """
    Chuyển đổi tất cả file MD trong folder
    
    Args:
        input_folder: Folder chứa file MD
        output_folder: Folder output
        pattern: Pattern filter
    
    Returns:
        List đường dẫn files đã tạo
    
    Example:
        >>> convert_folder(
        ...     'Segmentation/phase3_content/',
        ...     'Segmentation/phase4_rendered/'
        ... )
    """
    input_path = Path(input_folder)
    output_path = Path(output_folder)
    output_path.mkdir(parents=True, exist_ok=True)
    
    files = sorted(input_path.glob(pattern))
    results = []
    
    for md_file in files:
        docx_name = md_file.stem + '.docx'
        docx_path = output_path / docx_name
        
        result = convert_md_to_docx(str(md_file), str(docx_path))
        results.append(result)
    
    print(f"\n✅ Converted {len(results)} files")
    return results


if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description="Convert Markdown to DOCX")
    parser.add_argument('--file', '-f', help='Single MD file to convert')
    parser.add_argument('--folder', '-d', help='Folder containing MD files')
    parser.add_argument('--output', '-o', help='Output path or folder')
    parser.add_argument('--template', '-t', default='thesis', 
                        choices=['thesis', 'report', 'official'],
                        help='Template type')
    
    args = parser.parse_args()
    
    if args.file:
        convert_md_to_docx(args.file, args.output, args.template)
    elif args.folder:
        if not args.output:
            print("Please specify --output folder")
        else:
            convert_folder(args.folder, args.output)
    else:
        print("Please specify --file or --folder")
        parser.print_help()
