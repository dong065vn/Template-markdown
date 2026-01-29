"""
Thesis/Dissertation Template
=============================
Template cho luận văn, luận án theo chuẩn đại học Việt Nam
Tuân thủ Nghị định 30/2020/NĐ-CP
"""

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Dict, List, Optional
import os

# Import base styles
import sys
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from styles.base_styles import (
    STYLES, FONT_NAME, PageMargins, create_nd30_styles, apply_style
)


class ThesisTemplate:
    """Template generator cho luận văn/luận án"""
    
    def __init__(self, config: Dict = None):
        self.config = config or {}
        self.document = Document()
        self._setup_document()
    
    def _setup_document(self):
        """Thiết lập document cơ bản"""
        # Apply margins
        margins = PageMargins()
        for section in self.document.sections:
            margins.apply_to_section(section)
            section.page_width = Cm(21.0)   # A4
            section.page_height = Cm(29.7)  # A4
        
        # Create styles
        create_nd30_styles(self.document)
        
        # Set default font
        style = self.document.styles['Normal']
        font = style.font
        font.name = FONT_NAME
        font.size = Pt(14)
        
        # Set Vietnamese font for Asian text
        rFonts = style.element.rPr.rFonts
        rFonts.set(qn('w:eastAsia'), FONT_NAME)
    
    def add_cover_page(self, 
                       university: str,
                       faculty: str,
                       title: str,
                       author: str,
                       supervisor: str,
                       year: str,
                       thesis_type: str = "LUẬN VĂN TỐT NGHIỆP"):
        """Tạo trang bìa luận văn"""
        doc = self.document
        
        # University name
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(university.upper())
        run.font.name = FONT_NAME
        run.font.size = Pt(14)
        run.font.bold = True
        
        # Faculty
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(faculty.upper())
        run.font.name = FONT_NAME
        run.font.size = Pt(14)
        run.font.bold = True
        
        # Spacing
        for _ in range(3):
            doc.add_paragraph()
        
        # Thesis type
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(thesis_type)
        run.font.name = FONT_NAME
        run.font.size = Pt(16)
        run.font.bold = True
        
        # Spacing
        doc.add_paragraph()
        
        # Title
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title.upper())
        run.font.name = FONT_NAME
        run.font.size = Pt(18)
        run.font.bold = True
        
        # Spacing
        for _ in range(5):
            doc.add_paragraph()
        
        # Author info
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Sinh viên thực hiện: {author}")
        run.font.name = FONT_NAME
        run.font.size = Pt(14)
        
        # Supervisor
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Giảng viên hướng dẫn: {supervisor}")
        run.font.name = FONT_NAME
        run.font.size = Pt(14)
        
        # Spacing
        for _ in range(3):
            doc.add_paragraph()
        
        # Year
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(year)
        run.font.name = FONT_NAME
        run.font.size = Pt(14)
        run.font.bold = True
        
        # Page break
        doc.add_page_break()
        
        return self
    
    def add_declaration_page(self, author: str):
        """Trang lời cam đoan"""
        doc = self.document
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("LỜI CAM ĐOAN")
        run.font.name = FONT_NAME
        run.font.size = Pt(14)
        run.font.bold = True
        
        doc.add_paragraph()
        
        content = f"""Tôi xin cam đoan đây là công trình nghiên cứu của riêng tôi. Các số liệu, kết quả nêu trong luận văn là trung thực và chưa từng được ai công bố trong bất kỳ công trình nào khác.

Tác giả luận văn


{author}"""
        
        p = doc.add_paragraph(content)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(14)
        p.paragraph_format.first_line_indent = Cm(1.27)
        
        doc.add_page_break()
        return self
    
    def add_acknowledgment_page(self, content: str = ""):
        """Trang lời cảm ơn"""
        doc = self.document
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("LỜI CẢM ƠN")
        run.font.name = FONT_NAME
        run.font.size = Pt(14)
        run.font.bold = True
        
        doc.add_paragraph()
        
        p = doc.add_paragraph(content or "[Nội dung lời cảm ơn]")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(14)
        p.paragraph_format.first_line_indent = Cm(1.27)
        
        doc.add_page_break()
        return self
    
    def add_table_of_contents(self):
        """Thêm mục lục (placeholder - sẽ update bằng Word)"""
        doc = self.document
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("MỤC LỤC")
        run.font.name = FONT_NAME
        run.font.size = Pt(14)
        run.font.bold = True
        
        doc.add_paragraph()
        doc.add_paragraph("[Mục lục sẽ được tự động tạo sau khi hoàn thành nội dung]")
        
        doc.add_page_break()
        return self
    
    def add_chapter(self, chapter_num: int, title: str, content: str = ""):
        """Thêm chương mới"""
        doc = self.document
        
        # Chapter heading
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"CHƯƠNG {chapter_num}")
        run.font.name = FONT_NAME
        run.font.size = Pt(14)
        run.font.bold = True
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title.upper())
        run.font.name = FONT_NAME
        run.font.size = Pt(14)
        run.font.bold = True
        
        doc.add_paragraph()
        
        if content:
            p = doc.add_paragraph(content)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p.runs:
                run.font.name = FONT_NAME
                run.font.size = Pt(14)
            p.paragraph_format.first_line_indent = Cm(1.27)
        
        return self
    
    def add_section(self, section_num: str, title: str, content: str = ""):
        """Thêm mục (ví dụ: 1.1, 1.2)"""
        doc = self.document
        
        p = doc.add_paragraph()
        run = p.add_run(f"{section_num}. {title}")
        run.font.name = FONT_NAME
        run.font.size = Pt(14)
        run.font.bold = True
        
        if content:
            p = doc.add_paragraph(content)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p.runs:
                run.font.name = FONT_NAME
                run.font.size = Pt(14)
            p.paragraph_format.first_line_indent = Cm(1.27)
        
        return self
    
    def add_references(self, references: List[str] = None):
        """Thêm tài liệu tham khảo"""
        doc = self.document
        
        doc.add_page_break()
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("TÀI LIỆU THAM KHẢO")
        run.font.name = FONT_NAME
        run.font.size = Pt(14)
        run.font.bold = True
        
        doc.add_paragraph()
        
        if references:
            for i, ref in enumerate(references, 1):
                p = doc.add_paragraph(f"[{i}] {ref}")
                for run in p.runs:
                    run.font.name = FONT_NAME
                    run.font.size = Pt(14)
        
        return self
    
    def save(self, filepath: str):
        """Lưu document"""
        self.document.save(filepath)
        return filepath


def create_thesis_from_markdown(md_content: str, config: Dict) -> ThesisTemplate:
    """
    Tạo luận văn từ nội dung Markdown
    
    Args:
        md_content: Nội dung markdown thô từ AI
        config: Cấu hình (university, faculty, title, author, etc.)
    
    Returns:
        ThesisTemplate instance
    """
    template = ThesisTemplate(config)
    
    # Add cover page
    template.add_cover_page(
        university=config.get("university", "TRƯỜNG ĐẠI HỌC"),
        faculty=config.get("faculty", "KHOA"),
        title=config.get("title", "TÊN LUẬN VĂN"),
        author=config.get("author", "TÁC GIẢ"),
        supervisor=config.get("supervisor", "GIẢNG VIÊN HƯỚNG DẪN"),
        year=config.get("year", "2026"),
        thesis_type=config.get("thesis_type", "LUẬN VĂN TỐT NGHIỆP")
    )
    
    # Add declaration
    template.add_declaration_page(config.get("author", "TÁC GIẢ"))
    
    # Add acknowledgment
    template.add_acknowledgment_page()
    
    # Add TOC
    template.add_table_of_contents()
    
    # Parse markdown and add chapters
    # TODO: Implement markdown parser
    
    return template


if __name__ == "__main__":
    # Test creating a thesis
    config = {
        "university": "TRƯỜNG ĐẠI HỌC GIAO THÔNG VẬN TẢI",
        "faculty": "KHOA CÔNG NGHỆ THÔNG TIN",
        "title": "Nghiên cứu ứng dụng trí tuệ nhân tạo trong xử lý văn bản",
        "author": "Nguyễn Văn A",
        "supervisor": "PGS.TS. Trần Văn B",
        "year": "2026",
    }
    
    template = ThesisTemplate(config)
    template.add_cover_page(**config)
    template.add_declaration_page(config["author"])
    template.add_acknowledgment_page("Lời cảm ơn mẫu...")
    template.add_table_of_contents()
    template.add_chapter(1, "TỔNG QUAN", "Nội dung chương 1...")
    template.add_section("1.1", "Đặt vấn đề", "Nội dung mục 1.1...")
    template.add_chapter(2, "CƠ SỞ LÝ THUYẾT")
    template.add_references(["Tài liệu 1", "Tài liệu 2"])
    
    template.save("test_thesis.docx")
    print("Created test_thesis.docx")
