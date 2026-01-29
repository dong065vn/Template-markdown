"""
Official Document Templates
============================
Templates cho văn bản hành chính theo Nghị định 30/2020/NĐ-CP:
- Công văn
- Quyết định
- Tờ trình
- Thông báo
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from typing import Dict, List
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from styles.base_styles import FONT_NAME, PageMargins, create_nd30_styles


class OfficialDocumentTemplate:
    """Base template cho văn bản hành chính"""
    
    QUOC_HIEU = "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM"
    TIEU_NGU = "Độc lập - Tự do - Hạnh phúc"
    
    def __init__(self):
        self.document = Document()
        self._setup_document()
    
    def _setup_document(self):
        margins = PageMargins()
        for section in self.document.sections:
            margins.apply_to_section(section)
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)
        
        create_nd30_styles(self.document)
        
        style = self.document.styles['Normal']
        font = style.font
        font.name = FONT_NAME
        font.size = Pt(14)
        rFonts = style.element.rPr.rFonts
        rFonts.set(qn('w:eastAsia'), FONT_NAME)
    
    def add_header(self, co_quan_chu_quan: str, co_quan_ban_hanh: str,
                   so_ky_hieu: str, dia_danh: str, ngay_thang: str):
        """Thêm header văn bản hành chính (2 cột)"""
        doc = self.document
        
        # Tạo bảng 2 cột cho header
        table = doc.add_table(rows=4, cols=2)
        table.autofit = False
        
        # Set column widths
        for row in table.rows:
            row.cells[0].width = Cm(8)
            row.cells[1].width = Cm(9)
        
        # Cột trái: Cơ quan
        # Row 0: Cơ quan chủ quản
        cell = table.cell(0, 0)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(co_quan_chu_quan.upper())
        run.font.name = FONT_NAME
        run.font.size = Pt(12)
        
        # Row 1: Cơ quan ban hành
        cell = table.cell(1, 0)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(co_quan_ban_hanh.upper())
        run.font.name = FONT_NAME
        run.font.size = Pt(12)
        run.font.bold = True
        
        # Row 2: Gạch ngang
        cell = table.cell(2, 0)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("_______________")
        run.font.name = FONT_NAME
        run.font.size = Pt(12)
        
        # Row 3: Số, ký hiệu
        cell = table.cell(3, 0)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Số: {so_ky_hieu}")
        run.font.name = FONT_NAME
        run.font.size = Pt(13)
        
        # Cột phải: Quốc hiệu, tiêu ngữ
        # Row 0: Quốc hiệu
        cell = table.cell(0, 1)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(self.QUOC_HIEU)
        run.font.name = FONT_NAME
        run.font.size = Pt(12)
        run.font.bold = True
        
        # Row 1: Tiêu ngữ
        cell = table.cell(1, 1)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(self.TIEU_NGU)
        run.font.name = FONT_NAME
        run.font.size = Pt(13)
        run.font.bold = True
        
        # Row 2: Gạch ngang
        cell = table.cell(2, 1)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("_______________________________________")
        run.font.name = FONT_NAME
        run.font.size = Pt(12)
        
        # Row 3: Địa danh, ngày tháng
        cell = table.cell(3, 1)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{dia_danh}, {ngay_thang}")
        run.font.name = FONT_NAME
        run.font.size = Pt(13)
        run.font.italic = True
        
        doc.add_paragraph()
        return self
    
    def add_signature_block(self, chuc_vu: str, ho_ten: str, 
                           quyen_han: str = ""):
        """Thêm khối chữ ký"""
        doc = self.document
        
        # Tạo bảng 2 cột
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        
        # Cột phải: Chữ ký
        cell = table.cell(0, 1)
        cell.width = Cm(8)
        
        # Quyền hạn (nếu có)
        if quyen_han:
            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(quyen_han.upper())
            run.font.name = FONT_NAME
            run.font.size = Pt(14)
            run.font.bold = True
        
        # Chức vụ
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(chuc_vu.upper())
        run.font.name = FONT_NAME
        run.font.size = Pt(14)
        run.font.bold = True
        
        # Khoảng trống cho chữ ký
        for _ in range(3):
            cell.add_paragraph()
        
        # Họ tên
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(ho_ten)
        run.font.name = FONT_NAME
        run.font.size = Pt(14)
        run.font.bold = True
        
        return self
    
    def add_noi_nhan(self, noi_nhan_list: List[str]):
        """Thêm nơi nhận"""
        doc = self.document
        
        p = doc.add_paragraph()
        run = p.add_run("Nơi nhận:")
        run.font.name = FONT_NAME
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.italic = True
        
        for item in noi_nhan_list:
            p = doc.add_paragraph()
            run = p.add_run(f"- {item};")
            run.font.name = FONT_NAME
            run.font.size = Pt(11)
        
        return self
    
    def save(self, filepath: str):
        self.document.save(filepath)
        return filepath


class CongVanTemplate(OfficialDocumentTemplate):
    """Template cho Công văn"""
    
    def add_content(self, kinh_gui: str, noi_dung: str, trich_yeu: str = ""):
        doc = self.document
        
        # Trích yếu (V/v)
        if trich_yeu:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"V/v {trich_yeu}")
            run.font.name = FONT_NAME
            run.font.size = Pt(12)
        
        doc.add_paragraph()
        
        # Kính gửi
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Kính gửi: {kinh_gui}")
        run.font.name = FONT_NAME
        run.font.size = Pt(14)
        
        doc.add_paragraph()
        
        # Nội dung
        p = doc.add_paragraph(noi_dung)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(14)
        p.paragraph_format.first_line_indent = Cm(1.27)
        
        return self


class QuyetDinhTemplate(OfficialDocumentTemplate):
    """Template cho Quyết định"""
    
    def add_content(self, trich_yeu: str, can_cu: List[str], 
                    dieu_list: List[Dict], nguoi_ky: str):
        doc = self.document
        
        # Tên văn bản
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("QUYẾT ĐỊNH")
        run.font.name = FONT_NAME
        run.font.size = Pt(14)
        run.font.bold = True
        
        # Trích yếu
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(trich_yeu)
        run.font.name = FONT_NAME
        run.font.size = Pt(14)
        run.font.bold = True
        
        # Gạch ngang
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("_______________")
        run.font.name = FONT_NAME
        run.font.size = Pt(12)
        
        doc.add_paragraph()
        
        # Thẩm quyền
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(nguoi_ky.upper())
        run.font.name = FONT_NAME
        run.font.size = Pt(14)
        run.font.bold = True
        
        doc.add_paragraph()
        
        # Căn cứ
        for can_cu_item in can_cu:
            p = doc.add_paragraph()
            run = p.add_run(f"Căn cứ {can_cu_item};")
            run.font.name = FONT_NAME
            run.font.size = Pt(14)
            run.font.italic = True
            p.paragraph_format.first_line_indent = Cm(1.27)
        
        doc.add_paragraph()
        
        # QUYẾT ĐỊNH
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("QUYẾT ĐỊNH:")
        run.font.name = FONT_NAME
        run.font.size = Pt(14)
        run.font.bold = True
        
        doc.add_paragraph()
        
        # Các điều
        for dieu in dieu_list:
            p = doc.add_paragraph()
            run = p.add_run(f"Điều {dieu['so']}. {dieu['noi_dung']}")
            run.font.name = FONT_NAME
            run.font.size = Pt(14)
            p.paragraph_format.first_line_indent = Cm(1.27)
        
        return self


class ToTrinhTemplate(OfficialDocumentTemplate):
    """Template cho Tờ trình"""
    
    def add_content(self, kinh_trinh: str, trich_yeu: str, 
                    ly_do: str, de_xuat: str):
        doc = self.document
        
        # Tên văn bản
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("TỜ TRÌNH")
        run.font.name = FONT_NAME
        run.font.size = Pt(14)
        run.font.bold = True
        
        # Trích yếu
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(trich_yeu)
        run.font.name = FONT_NAME
        run.font.size = Pt(14)
        run.font.bold = True
        
        # Gạch ngang
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("_______________")
        
        doc.add_paragraph()
        
        # Kính trình
        p = doc.add_paragraph()
        run = p.add_run(f"Kính trình: {kinh_trinh}")
        run.font.name = FONT_NAME
        run.font.size = Pt(14)
        run.font.bold = True
        
        doc.add_paragraph()
        
        # Phần I: Lý do
        p = doc.add_paragraph()
        run = p.add_run("I. LÝ DO, SỰ CẦN THIẾT")
        run.font.name = FONT_NAME
        run.font.size = Pt(14)
        run.font.bold = True
        
        p = doc.add_paragraph(ly_do)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(14)
        p.paragraph_format.first_line_indent = Cm(1.27)
        
        doc.add_paragraph()
        
        # Phần II: Đề xuất
        p = doc.add_paragraph()
        run = p.add_run("II. NỘI DUNG ĐỀ XUẤT")
        run.font.name = FONT_NAME
        run.font.size = Pt(14)
        run.font.bold = True
        
        p = doc.add_paragraph(de_xuat)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(14)
        p.paragraph_format.first_line_indent = Cm(1.27)
        
        return self


if __name__ == "__main__":
    # Test Công văn
    cv = CongVanTemplate()
    cv.add_header(
        co_quan_chu_quan="BỘ GIÁO DỤC VÀ ĐÀO TẠO",
        co_quan_ban_hanh="TRƯỜNG ĐẠI HỌC ABC",
        so_ky_hieu="123/CV-ĐHABC",
        dia_danh="Hà Nội",
        ngay_thang="ngày 29 tháng 01 năm 2026"
    )
    cv.add_content(
        kinh_gui="Ban Giám hiệu Trường",
        noi_dung="Nội dung công văn...",
        trich_yeu="Đề nghị phê duyệt kế hoạch"
    )
    cv.add_signature_block("HIỆU TRƯỞNG", "Nguyễn Văn A", "KT.")
    cv.add_noi_nhan(["Như trên", "Lưu: VT, TCHC"])
    cv.save("test_congvan.docx")
    print("Created test_congvan.docx")
