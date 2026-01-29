"""
Report Template
================
Template cho báo cáo (báo cáo dự án, báo cáo thực tập, báo cáo kỹ thuật)
Tuân thủ Nghị định 30/2020/NĐ-CP
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from typing import Dict, List
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from styles.base_styles import FONT_NAME, PageMargins, create_nd30_styles


class ReportTemplate:
    """Template generator cho báo cáo"""
    
    def __init__(self, config: Dict = None):
        self.config = config or {}
        self.document = Document()
        self._setup_document()
    
    def _setup_document(self):
        """Thiết lập document cơ bản"""
        margins = PageMargins()
        for section in self.document.sections:
            margins.apply_to_section(section)
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)
        
        create_nd30_styles(self.document)
        
        style = self.document.styles['Normal']
        font = style.font
        font.name = FONT_NAME
        font.size = Pt(14)
        
        rFonts = style.element.rPr.rFonts
        rFonts.set(qn('w:eastAsia'), FONT_NAME)
    
    def add_cover_page(self,
                       organization: str,
                       department: str,
                       title: str,
                       author: str,
                       date: str,
                       report_type: str = "BÁO CÁO"):
        """Tạo trang bìa báo cáo"""
        doc = self.document
        
        # Organization
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(organization.upper())
        run.font.name = FONT_NAME
        run.font.size = Pt(14)
        run.font.bold = True
        
        # Department
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(department.upper())
        run.font.name = FONT_NAME
        run.font.size = Pt(14)
        run.font.bold = True
        
        # Spacing
        for _ in range(4):
            doc.add_paragraph()
        
        # Report type
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(report_type)
        run.font.name = FONT_NAME
        run.font.size = Pt(16)
        run.font.bold = True
        
        doc.add_paragraph()
        
        # Title
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title.upper())
        run.font.name = FONT_NAME
        run.font.size = Pt(18)
        run.font.bold = True
        
        # Spacing
        for _ in range(6):
            doc.add_paragraph()
        
        # Author
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Người thực hiện: {author}")
        run.font.name = FONT_NAME
        run.font.size = Pt(14)
        
        # Spacing
        for _ in range(3):
            doc.add_paragraph()
        
        # Date
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(date)
        run.font.name = FONT_NAME
        run.font.size = Pt(14)
        run.font.italic = True
        
        doc.add_page_break()
        return self
    
    def add_table_of_contents(self):
        """Thêm mục lục"""
        doc = self.document
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("MỤC LỤC")
        run.font.name = FONT_NAME
        run.font.size = Pt(14)
        run.font.bold = True
        
        doc.add_paragraph()
        doc.add_paragraph("[Mục lục sẽ được tự động tạo]")
        doc.add_page_break()
        return self
    
    def add_section(self, section_num: str, title: str, content: str = ""):
        """Thêm phần/mục"""
        doc = self.document
        
        p = doc.add_paragraph()
        run = p.add_run(f"{section_num}. {title.upper()}")
        run.font.name = FONT_NAME
        run.font.size = Pt(14)
        run.font.bold = True
        
        if content:
            p = doc.add_paragraph(content)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p.runs:
                run.font.name = FONT_NAME
                run.font.size = Pt(14)
            p.paragraph_format.first_line_indent = Cm(1.27)
        
        return self
    
    def add_subsection(self, subsection_num: str, title: str, content: str = ""):
        """Thêm mục con"""
        doc = self.document
        
        p = doc.add_paragraph()
        run = p.add_run(f"{subsection_num}. {title}")
        run.font.name = FONT_NAME
        run.font.size = Pt(14)
        run.font.bold = True
        
        if content:
            p = doc.add_paragraph(content)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p.runs:
                run.font.name = FONT_NAME
                run.font.size = Pt(14)
            p.paragraph_format.first_line_indent = Cm(1.27)
        
        return self
    
    def add_conclusion(self, content: str = ""):
        """Thêm kết luận"""
        doc = self.document
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("KẾT LUẬN")
        run.font.name = FONT_NAME
        run.font.size = Pt(14)
        run.font.bold = True
        
        doc.add_paragraph()
        
        if content:
            p = doc.add_paragraph(content)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p.runs:
                run.font.name = FONT_NAME
                run.font.size = Pt(14)
            p.paragraph_format.first_line_indent = Cm(1.27)
        
        return self
    
    def add_appendix(self, appendix_num: str, title: str, content: str = ""):
        """Thêm phụ lục"""
        doc = self.document
        
        doc.add_page_break()
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"PHỤ LỤC {appendix_num}")
        run.font.name = FONT_NAME
        run.font.size = Pt(14)
        run.font.bold = True
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title.upper())
        run.font.name = FONT_NAME
        run.font.size = Pt(14)
        run.font.bold = True
        
        if content:
            doc.add_paragraph()
            p = doc.add_paragraph(content)
            for run in p.runs:
                run.font.name = FONT_NAME
                run.font.size = Pt(14)
        
        return self
    
    def save(self, filepath: str):
        """Lưu document"""
        self.document.save(filepath)
        return filepath


def create_report_from_markdown(md_content: str, config: Dict) -> ReportTemplate:
    """Tạo báo cáo từ nội dung Markdown"""
    template = ReportTemplate(config)
    
    template.add_cover_page(
        organization=config.get("organization", "TỔ CHỨC"),
        department=config.get("department", "PHÒNG BAN"),
        title=config.get("title", "TÊN BÁO CÁO"),
        author=config.get("author", "TÁC GIẢ"),
        date=config.get("date", "Ngày ... tháng ... năm ..."),
        report_type=config.get("report_type", "BÁO CÁO")
    )
    
    template.add_table_of_contents()
    
    return template


if __name__ == "__main__":
    config = {
        "organization": "CÔNG TY ABC",
        "department": "PHÒNG CÔNG NGHỆ THÔNG TIN",
        "title": "Báo cáo tiến độ dự án XYZ",
        "author": "Nguyễn Văn A",
        "date": "Tháng 01/2026",
    }
    
    template = ReportTemplate(config)
    template.add_cover_page(**config)
    template.add_table_of_contents()
    template.add_section("I", "TỔNG QUAN", "Nội dung tổng quan...")
    template.add_subsection("1.1", "Mục tiêu", "Nội dung mục tiêu...")
    template.add_section("II", "NỘI DUNG THỰC HIỆN")
    template.add_conclusion("Kết luận báo cáo...")
    template.add_appendix("A", "BẢNG BIỂU")
    
    template.save("test_report.docx")
    print("Created test_report.docx")
