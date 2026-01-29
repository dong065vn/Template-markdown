"""
DOCX Parser
============
Parse DOCX files using python-docx to IR
"""

import os
from pathlib import Path
from typing import List, Optional

from docx import Document as DocxDocument
from docx.shared import Pt

from .ir import (
    Document, Section, Block, BlockType, Run, ListItem, TableRow, TableCell,
    create_paragraph, create_heading
)


class DOCXParser:
    """Parse DOCX files to Intermediate Representation"""
    
    def __init__(self):
        pass
    
    def parse(self, docx_path: str) -> Document:
        """
        Parse DOCX file to Document IR
        
        Args:
            docx_path: Đường dẫn đến file DOCX
        
        Returns:
            Document object chứa nội dung đã parse
        """
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"File not found: {docx_path}")
        
        docx = DocxDocument(docx_path)
        
        # Create IR Document
        ir_doc = Document(
            title=Path(docx_path).stem,
            source_path=docx_path,
            source_type="docx"
        )
        
        # Get core properties if available
        try:
            if docx.core_properties.title:
                ir_doc.title = docx.core_properties.title
            if docx.core_properties.author:
                ir_doc.author = docx.core_properties.author
        except:
            pass
        
        current_section = None
        
        for para in docx.paragraphs:
            text = para.text.strip()
            
            if not text:
                continue
            
            # Detect heading by style
            style_name = para.style.name.lower() if para.style else ""
            
            if "heading 1" in style_name or style_name == "title":
                # H1 - Start new section
                if current_section:
                    ir_doc.sections.append(current_section)
                current_section = Section(title=text, level=1)
            
            elif "heading 2" in style_name:
                if current_section is None:
                    current_section = Section(title="Document", level=1)
                current_section.blocks.append(create_heading(text, 2))
            
            elif "heading 3" in style_name:
                if current_section is None:
                    current_section = Section(title="Document", level=1)
                current_section.blocks.append(create_heading(text, 3))
            
            elif "heading 4" in style_name:
                if current_section is None:
                    current_section = Section(title="Document", level=1)
                current_section.blocks.append(create_heading(text, 4))
            
            elif "list" in style_name:
                # List item
                if current_section is None:
                    current_section = Section(title="Document", level=1)
                
                # Check if last block is a list
                if current_section.blocks and current_section.blocks[-1].type == BlockType.LIST:
                    # Add to existing list
                    current_section.blocks[-1].items.append(
                        ListItem(content=self._parse_runs(para))
                    )
                else:
                    # Create new list
                    is_ordered = "number" in style_name
                    current_section.blocks.append(Block(
                        type=BlockType.LIST,
                        ordered=is_ordered,
                        items=[ListItem(content=self._parse_runs(para))]
                    ))
            
            else:
                # Normal paragraph
                if current_section is None:
                    current_section = Section(title="Document", level=1)
                
                runs = self._parse_runs(para)
                current_section.blocks.append(Block(
                    type=BlockType.PARAGRAPH,
                    runs=runs
                ))
        
        # Parse tables
        for table in docx.tables:
            if current_section is None:
                current_section = Section(title="Document", level=1)
            
            table_block = self._parse_table(table)
            current_section.blocks.append(table_block)
        
        # Add last section
        if current_section:
            ir_doc.sections.append(current_section)
        
        return ir_doc
    
    def _parse_runs(self, para) -> List[Run]:
        """Parse runs từ paragraph"""
        runs = []
        for run in para.runs:
            if run.text:
                runs.append(Run(
                    text=run.text,
                    bold=run.bold or False,
                    italic=run.italic or False,
                    underline=run.underline is not None,
                    font_name=run.font.name,
                    font_size=run.font.size.pt if run.font.size else None
                ))
        return runs
    
    def _parse_table(self, table) -> Block:
        """Parse table to Block"""
        rows = []
        
        for i, row in enumerate(table.rows):
            cells = []
            for cell in row.cells:
                text = cell.text.strip()
                cells.append(TableCell(
                    content=[Run(text=text)]
                ))
            rows.append(TableRow(
                cells=cells,
                is_header=(i == 0)  # First row is header
            ))
        
        return Block(
            type=BlockType.TABLE,
            rows=rows
        )


def parse_docx(docx_path: str) -> Document:
    """
    Hàm tiện ích để parse DOCX
    
    Args:
        docx_path: Đường dẫn file DOCX
    
    Returns:
        Document IR
    
    Example:
        >>> doc = parse_docx("thesis.docx")
        >>> print(doc.get_section_count())
    """
    parser = DOCXParser()
    return parser.parse(docx_path)


if __name__ == "__main__":
    import sys
    
    if len(sys.argv) > 1:
        doc = parse_docx(sys.argv[1])
        print(f"Title: {doc.title}")
        print(f"Author: {doc.author}")
        print(f"Sections: {doc.get_section_count()}")
        print(f"Blocks: {doc.get_total_blocks()}")
        print("\n--- Markdown ---\n")
        print(doc.to_markdown()[:2000])
    else:
        print("Usage: python docx_parser.py <docx_file>")
